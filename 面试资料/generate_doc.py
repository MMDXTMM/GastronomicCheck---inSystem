# -*- coding: utf-8 -*-
"""生成 hm-dianping 面试复习 Word 文档。运行：python generate_doc.py（建议 Windows 下先设 PYTHONUTF8=1）"""
import sys
import io
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "hm-dianping面试复习手册.docx")


def set_cn_font(run, name="微软雅黑", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_title(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_cn_font(run, size=18 if level == 1 else 14, bold=True)
    return p


def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_cn_font(run, bold=bold)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_cn_font(run)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_cn_font(run, bold=True, size=10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_cn_font(run, size=10)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def build():
    doc = Document()
    # 默认正文
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)

    # ===== 封面 =====
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("hm-dianping 项目\n面试复习手册")
    set_cn_font(r, size=22, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("（模块设计 · Spring 注解 · 高频面试题）")
    set_cn_font(r2, size=12)

    doc.add_paragraph()
    add_para(doc, "项目名称：hm-dianping（黑马点评）")
    add_para(doc, "技术栈：Spring Boot 2.3.12 + MyBatis-Plus + MySQL + Redis + Redisson")
    add_para(doc, "架构模式：Spring Boot 单体应用，经典三层架构（Controller → Service → Mapper）")
    doc.add_page_break()

    # ===== 第一章 =====
    add_title(doc, "第一章  项目整体介绍", 1)
    add_para(doc, "这是一个类似「大众点评」的本地生活平台后端项目。虽然是单体应用，没有拆成多个 Maven 子模块，但按业务域划分了清晰的模块，每个模块对应一套 Controller / Service / Mapper / Entity。")

    add_title(doc, "1.1  项目目录结构", 2)
    add_code_block(doc, """hm-dianping/
├── src/main/java/com/hmdp/
│   ├── HmDianPingApplication.java   # 启动类
│   ├── controller/                  # 接口层（9 个 Controller）
│   ├── service/ + service/impl/     # 业务层
│   ├── mapper/                      # 数据访问层
│   ├── entity/                      # 数据库实体
│   ├── dto/                         # 数据传输对象
│   ├── config/                      # Spring 配置
│   └── utils/                       # 工具类（缓存、锁、拦截器等）
└── src/main/resources/
    ├── application.yaml
    ├── db/hmdp.sql
    └── mapper/VoucherMapper.xml""")

    add_title(doc, "1.2  分层架构", 2)
    add_code_block(doc, """客户端
   ↓
Controller（参数校验、调用 Service、返回 Result）
   ↓
Service（业务逻辑、事务、缓存、分布式锁）
   ↓
Mapper（MyBatis-Plus / 自定义 SQL）
   ↓
MySQL / Redis""")

    add_title(doc, "1.3  30 秒开场白（可直接背诵）", 2)
    add_para(doc, "「我做的是一个类似大众点评的本地生活平台，后端是 Spring Boot 单体应用，采用经典三层架构：Controller 接口层、Service 业务层、Mapper 数据访问层。技术栈是 Spring Boot + MyBatis-Plus + MySQL + Redis + Redisson。项目按业务域划分模块，每个模块一套 Controller/Service/Mapper/Entity，职责清晰。亮点主要在店铺 Redis 缓存（穿透/击穿/一致性）和秒杀模块（分布式锁 + 乐观扣库存防超卖）。」")

    doc.add_page_break()

    # ===== 第二章 业务模块 =====
    add_title(doc, "第二章  业务模块详解", 1)

    add_table(doc,
        ["模块", "主要类", "API 前缀", "核心功能"],
        [
            ["用户模块", "UserController / UserServiceImpl", "/user", "短信验证码登录、登出、用户信息"],
            ["店铺模块", "ShopController / ShopServiceImpl", "/shop", "店铺 CRUD、按类型/名称查询"],
            ["店铺类型", "ShopTypeController", "/shop-type", "店铺分类列表"],
            ["优惠券模块", "VoucherController / VoucherServiceImpl", "/voucher", "普通券、秒杀券创建与查询"],
            ["秒杀/库存/订单", "VoucherOrderController / VoucherOrderServiceImpl", "/voucher-order", "秒杀下单、库存扣减、防重复下单"],
            ["探店笔记", "BlogController / BlogServiceImpl", "/blog", "发布笔记、点赞、Feed 流"],
            ["博客评论", "BlogCommentsController", "/blog-comments", "评论功能"],
            ["关注模块", "FollowController / FollowServiceImpl", "/follow", "关注/取关、共同关注"],
            ["文件上传", "UploadController", "/upload", "博客图片上传/删除"],
            ["基础设施", "config/ + utils/", "—", "Redis 缓存、分布式锁、登录拦截、全局异常"],
        ])

    # 各模块详细
    modules_detail = [
        ("2.1  用户模块", [
            "功能：手机号 + 验证码登录、登出、查询当前用户/用户信息",
            "设计要点：",
            "  · 验证码存 Redis，带 TTL，不用 Session 存验证码",
            "  · 登录成功后生成 Token，存 Redis，前端请求头携带 Token",
            "  · 双拦截器：RefreshTokenInterceptor（Token 续期，order=0）+ LoginInterceptor（登录校验）",
            "  · 白名单路径免登录：/shop/**、/voucher/**、/user/code、/user/login 等",
            "主要 API：POST /user/code、POST /user/login、GET /user/me、POST /user/logout",
        ]),
        ("2.2  店铺模块（缓存亮点）", [
            "功能：店铺查询、新增、修改，按类型/名称分页查询",
            "设计要点（面试重点）：",
            "  · 店铺详情走 Redis 缓存，Key 规范：cache:shop:{id}",
            "  · 缓存穿透：缓存空值（CacheClient.queryWithPassThrough）",
            "  · 缓存击穿：互斥锁重建 或 逻辑过期（queryWithLogicalExpire）",
            "  · 缓存一致性：更新 DB 后删除缓存（Cache Aside 模式）",
            "  · 封装 CacheClient 工具类，统一缓存读写逻辑",
            "主要 API：GET /shop/{id}、POST /shop、PUT /shop、GET /shop/of/type",
        ]),
        ("2.3  优惠券 + 秒杀/库存模块（核心亮点）", [
            "功能：发布普通券/秒杀券、查询店铺优惠券、秒杀下单",
            "表设计：",
            "  · tb_voucher：优惠券基本信息",
            "  · tb_seckill_voucher：秒杀扩展信息（stock 库存、begin_time、end_time），与 voucher 一对一",
            "  · tb_voucher_order：秒杀订单",
            "秒杀下单流程：",
            "  1. 查询秒杀券，校验开始/结束时间",
            "  2. 判断 stock >= 1",
            "  3. Redisson 分布式锁（lock:order:{userId}）防同一用户并发重复下单",
            "  4. 通过 AOP 代理调用 createVoucherOrder()（保证 @Transactional 生效）",
            "  5. 校验一人一单（查订单表）",
            "  6. 乐观扣库存：UPDATE SET stock=stock-1 WHERE voucher_id=? AND stock>0",
            "  7. 创建订单，订单 ID 由 RedisIdWorker 生成",
            "主要 API：POST /voucher/seckill、POST /voucher-order/seckill/{id}",
        ]),
        ("2.4  探店笔记模块", [
            "功能：发笔记、热门笔记、点赞、个人笔记、关注的人的 Feed 流",
            "设计要点：",
            "  · 点赞用 Redis ZSet 存储（key: blog:liked:{blogId}，member: userId，score: 时间戳）",
            "  · Feed 流用 Redis ZSet 滚动分页（ScrollResult，按 score 排序）",
            "主要 API：POST /blog、GET /blog/hot、PUT /blog/like/{id}、GET /blog/of/follow",
        ]),
        ("2.5  关注模块", [
            "功能：关注/取关、判断是否关注、查共同关注",
            "设计：关系存 MySQL tb_follow；共同关注可用 Redis Set 交集优化",
            "主要 API：POST /follow/{id}/{isFollow}、GET /follow/or/not/{id}、GET /follow/common/{id}",
        ]),
        ("2.6  基础设施层", [
            "WebExceptionAdvice：@RestControllerAdvice 全局异常，统一返回 Result",
            "MybatisConfig：MyBatis-Plus 分页插件",
            "RedissonConfig：Redisson 客户端 Bean（分布式锁）",
            "MvcConfig：注册登录拦截器、配置白名单",
            "CacheClient：Redis 缓存工具（穿透/击穿解决方案）",
            "RedisIdWorker：基于 Redis 的全局唯一 ID 生成器",
            "SimpleRedisLock：简易 Redis 分布式锁（Lua 脚本释放锁）",
        ]),
    ]

    for title, lines in modules_detail:
        add_title(doc, title, 2)
        for line in lines:
            if line.startswith("  ·"):
                add_bullet(doc, line.strip("  · "))
            elif line.endswith("：") or line.endswith("要点：") or line.endswith("流程：") or line.endswith("设计："):
                add_para(doc, line, bold=True)
            else:
                add_para(doc, line)

    doc.add_page_break()

    # ===== 第三章 Spring 注解 =====
    add_title(doc, "第三章  Spring 注解速查手册", 1)
    add_para(doc, "本项目使用的 Spring 及相关框架注解汇总，按类别整理，便于面试快速回忆。")

    spring_sections = [
        ("3.1  启动与扫描", [
            ["@SpringBootApplication", "HmDianPingApplication", "启动 Spring Boot，开启自动配置"],
            ["@MapperScan(\"com.hmdp.mapper\")", "启动类", "扫描 MyBatis Mapper 接口，注册为 Bean"],
            ["@EnableAspectJAutoProxy(exposeProxy=true)", "启动类", "开启 AOP；exposeProxy 暴露代理对象，解决同类 @Transactional 自调用失效"],
        ]),
        ("3.2  组件注册（IoC 容器）", [
            ["@Service", "所有 *ServiceImpl", "标记业务层 Bean，由 Spring 管理"],
            ["@RestController", "所有 *Controller", "REST 控制器 = @Controller + @ResponseBody，返回值自动转 JSON"],
            ["@Component", "CacheClient、RedisIdWorker", "通用组件 Bean"],
            ["@Configuration", "MvcConfig、RedissonConfig、MybatisConfig", "标记配置类"],
            ["@Bean", "RedissonConfig、MybatisConfig", "在 @Configuration 类中向容器注册第三方 Bean"],
        ]),
        ("3.3  Web 层注解（Controller）", [
            ["@RequestMapping(\"/shop\")", "Controller 类", "类级别路由前缀"],
            ["@GetMapping(\"/{id}\")", "Controller 方法", "映射 HTTP GET 请求"],
            ["@PostMapping(\"/login\")", "Controller 方法", "映射 HTTP POST 请求"],
            ["@PutMapping(\"/like/{id}\")", "Controller 方法", "映射 HTTP PUT 请求"],
            ["@PathVariable(\"id\")", "方法参数", "绑定 URL 路径变量"],
            ["@RequestParam(\"phone\")", "方法参数", "绑定查询参数 / 表单参数"],
            ["@RequestBody Shop shop", "方法参数", "将请求体 JSON 反序列化为 Java 对象"],
        ]),
        ("3.4  依赖注入", [
            ["@Resource", "Controller/Service 字段", "按名称注入依赖（JSR-250 标准，本项目主要使用）"],
            ["构造器注入", "ShopServiceImpl", "通过构造函数注入 StringRedisTemplate（推荐方式）"],
        ]),
        ("3.5  事务与 AOP", [
            ["@Transactional", "VoucherOrderServiceImpl 等", "声明式事务；方法内所有 DB 操作在同一事务中，异常则回滚"],
        ]),
        ("3.6  全局异常处理", [
            ["@RestControllerAdvice", "WebExceptionAdvice", "全局 REST 异常处理器（= @ControllerAdvice + @ResponseBody）"],
            ["@ExceptionHandler(RuntimeException.class)", "异常处理方法", "捕获指定类型异常并返回统一 Result"],
        ]),
        ("3.7  MyBatis-Plus 注解（Entity 层）", [
            ["@TableName(\"tb_shop\")", "Entity 类", "指定对应的数据库表名"],
            ["@TableId(value=\"id\", type=IdType.AUTO)", "Entity 主键字段", "主键映射；AUTO=自增，INPUT=手动输入"],
            ["@TableField(exist=false)", "Entity 非表字段", "标记该字段不是数据库列（如 Voucher.stock）"],
        ]),
        ("3.8  其他常用注解", [
            ["@Slf4j", "UserController 等", "Lombok：自动生成 log 日志对象"],
            ["@Data", "Entity/DTO", "Lombok：生成 getter/setter/toString/equals/hashCode"],
            ["@Override", "ServiceImpl", "Java：标记重写接口/父类方法"],
            ["@Param(\"shopId\")", "VoucherMapper", "MyBatis：Mapper 方法参数绑定"],
            ["@JsonIgnore", "ShopType Entity", "Jackson：JSON 序列化时忽略该字段"],
        ]),
    ]

    for sec_title, rows in spring_sections:
        add_title(doc, sec_title, 2)
        add_table(doc, ["注解", "使用位置", "作用说明"], rows)

    add_title(doc, "3.9  各模块注解对照表", 2)
    add_table(doc,
        ["模块", "Controller 注解", "Service 注解", "特殊注解"],
        [
            ["用户", "@RestController @RequestMapping @PostMapping @GetMapping @RequestParam @RequestBody @Slf4j", "@Service @Resource @Slf4j", "—"],
            ["店铺", "同上 + @PathVariable", "@Service @Transactional @Resource", "构造器注入 RedisTemplate"],
            ["优惠券", "同上", "@Service @Transactional", "—"],
            ["秒杀/库存", "@RestController @PostMapping", "@Service @Transactional @Resource", "@EnableAspectJAutoProxy（启动类）"],
            ["博客", "同上 + @PutMapping", "@Service @Resource", "—"],
            ["关注", "同上", "@Service @Resource", "—"],
            ["配置类", "—", "—", "@Configuration @Bean"],
            ["异常处理", "@RestControllerAdvice @ExceptionHandler", "—", "—"],
        ])

    doc.add_page_break()

    # ===== 第四章 面试高频题 =====
    add_title(doc, "第四章  面试高频问答", 1)

    qa_list = [
        ("Q1：项目有哪些模块？怎么划分的？",
         "按业务域划分，非 Maven 多模块。每个域有 Controller-Service-Mapper-Entity 一套：用户、店铺、优惠券秒杀、探店笔记、关注、上传，以及 config/utils 基础设施。模块间通过 Service 接口调用，不跨模块直接操作 Mapper。"),
        ("Q2：为什么用单体而不是微服务？",
         "学习/练手项目，业务规模适中，单体 + 业务分包足够。后续可按用户、店铺、秒杀、内容等域拆微服务；秒杀可单独做高并发优化。"),
        ("Q3：模块之间怎么解耦？",
         "通过 Service 接口调用（如秒杀调 ISeckillVoucherService，不直接写 Mapper）。公共能力（缓存、锁、ID 生成）抽到 utils 和 config。"),
        ("Q4：秒杀怎么保证不超卖？",
         "三层保障：① Redisson 锁，同一用户并发只让一个进；② 一人一单，查订单表；③ 乐观扣库存 UPDATE stock=stock-1 WHERE stock>0，影响行数为 0 则失败；④ @Transactional 保证扣库存和建单原子性。"),
        ("Q5：为什么启动类要加 @EnableAspectJAutoProxy(exposeProxy=true)？",
         "seckillVoucher() 和 createVoucherOrder() 在同一个类。直接 this.createVoucherOrder() 不走代理，@Transactional 不生效。exposeProxy=true 后可通过 AopContext.currentProxy() 拿到代理对象再调用，事务才能生效。"),
        ("Q6：缓存怎么设计的？解决了哪些问题？",
         "Key 规范如 cache:shop:{id}。穿透：缓存空值；击穿：互斥锁重建或逻辑过期；一致性：更新 DB 后删缓存（Cache Aside）。封装 CacheClient 统一处理。"),
        ("Q7：登录怎么设计的？",
         "无状态 Token + Redis 存用户信息。双拦截器：RefreshTokenInterceptor 续期（order=0），LoginInterceptor 校验登录。店铺浏览、登录注册等在白名单免登录。"),
        ("Q8：库存存在哪？怎么扣减？",
         "库存字段 stock 在 tb_seckill_voucher 表。创建秒杀券时从 Voucher 传入写入 SeckillVoucher。扣减用 MyBatis-Plus：update().setSql(\"stock = stock - 1\").eq(\"voucher_id\", id).gt(\"stock\", 0)，数据库层面保证不超卖。"),
        ("Q9：订单 ID 怎么生成的？",
         "RedisIdWorker：Redis INCR + 时间戳 + 业务前缀，保证分布式环境下全局唯一，趋势递增。"),
        ("Q10：@Resource 和 @Autowired 区别？",
         "本项目主要用 @Resource（JSR-250）：默认按名称注入，找不到再按类型。@Autowired（Spring）：默认按类型注入，可配合 @Qualifier 指定名称。功能类似，@Resource 不依赖 Spring 特有 API。"),
        ("Q11：@RestController 和 @Controller 区别？",
         "@RestController = @Controller + @ResponseBody，方法返回值直接序列化为 JSON。@Controller 通常配合视图解析器返回页面；REST API 项目用 @RestController。"),
        ("Q12：@Transactional 什么时候会失效？",
         "① 同类内部自调用（未走代理）；② 方法非 public；③ 异常被 catch 未抛出；④ 数据库引擎不支持事务（如 MyISAM）。本项目通过 AopContext.currentProxy() 解决自调用问题。"),
    ]

    for q, a in qa_list:
        add_para(doc, q, bold=True)
        add_para(doc, a, indent=True)
        doc.add_paragraph()

    doc.add_page_break()

    # ===== 第五章 1分钟完整版 =====
    add_title(doc, "第五章  1 分钟完整口述稿", 1)
    add_para(doc, "「项目是一个本地生活点评平台，Spring Boot 单体，三层架构。业务上分用户、店铺、优惠券秒杀、探店笔记、关注等模块，每个模块都是 Controller-Service-Mapper-Entity 一套。」")
    add_para(doc, "「用户模块用 Redis 存验证码和 Token，双拦截器做登录校验和续期。」")
    add_para(doc, "「店铺模块重点做了 Redis 缓存，用 CacheClient 解决了穿透、击穿和一致性问题。」")
    add_para(doc, "「秒杀模块库存放秒杀券表，下单用 Redisson 锁防重复、SQL 乐观更新防超卖、@Transactional 保证扣库存和建单一致，订单号用 Redis 全局 ID 生成。启动类开了 @EnableAspectJAutoProxy 解决事务自调用。」")
    add_para(doc, "「博客模块用 Redis ZSet 做点赞和 Feed 流。全局 Result 统一返回，RestControllerAdvice 做异常处理。」")
    add_para(doc, "「整体按业务域划分、分层清晰，高并发和缓存亮点集中在店铺和秒杀两个模块。」")

    add_title(doc, "5.1  表达技巧", 2)
    add_bullet(doc, "先业务后技术：先说模块解决什么问题，再说 Redis、锁等技术细节")
    add_bullet(doc, "挑 2 个亮点深挖：建议店铺缓存 + 秒杀库存，不要每个模块都浅尝辄止")
    add_bullet(doc, "诚实说边界：如「Redis 预减库存 Key 已定义，当前以 DB 扣减为主，高并发可升级」")
    add_bullet(doc, "准备画简图：秒杀流程或三层架构，白板面试很加分")

    doc.add_page_break()

    # ===== 附录 API 清单 =====
    add_title(doc, "附录  主要 API 接口清单", 1)
    add_table(doc,
        ["模块", "方法", "路径", "说明"],
        [
            ["用户", "POST", "/user/code", "发送验证码"],
            ["用户", "POST", "/user/login", "登录"],
            ["用户", "GET", "/user/me", "当前用户信息"],
            ["用户", "POST", "/user/logout", "登出"],
            ["店铺", "GET", "/shop/{id}", "查询店铺详情"],
            ["店铺", "POST", "/shop", "新增店铺"],
            ["店铺", "PUT", "/shop", "修改店铺"],
            ["店铺", "GET", "/shop/of/type", "按类型分页查询"],
            ["优惠券", "POST", "/voucher", "新增普通券"],
            ["优惠券", "POST", "/voucher/seckill", "新增秒杀券"],
            ["优惠券", "GET", "/voucher/list/{shopId}", "查询店铺优惠券"],
            ["秒杀", "POST", "/voucher-order/seckill/{id}", "秒杀下单"],
            ["博客", "POST", "/blog", "发布笔记"],
            ["博客", "GET", "/blog/hot", "热门笔记"],
            ["博客", "PUT", "/blog/like/{id}", "点赞"],
            ["博客", "GET", "/blog/of/follow", "关注的人 Feed 流"],
            ["关注", "POST", "/follow/{id}/{isFollow}", "关注/取关"],
            ["关注", "GET", "/follow/or/not/{id}", "是否关注"],
            ["关注", "GET", "/follow/common/{id}", "共同关注"],
            ["上传", "POST", "/upload/blog", "上传图片"],
        ])

    add_para(doc, "")
    add_para(doc, "— 文档完 —", bold=True)
    add_para(doc, "生成说明：基于 hm-dianping 项目源码整理，供面试复习使用。")

    doc.save(OUTPUT)
    print(f"已生成: {OUTPUT}")


if __name__ == "__main__":
    build()
